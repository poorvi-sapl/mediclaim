"""Generate a one-page ClaimLens scoring reference (.docx) — plain black & white."""
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.5)
    s.left_margin = s.right_margin = Inches(0.6)

normal = doc.styles["Normal"].font
normal.name = "Calibri"; normal.size = Pt(10)


def heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(8); p.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)


# Title
t = doc.add_paragraph(); t.space_after = Pt(0)
r = t.add_run("ClaimLens — Fraud Risk Scoring"); r.bold = True; r.font.size = Pt(15)
sub = doc.add_paragraph(); sub.space_after = Pt(2)
rs = sub.add_run("How the 0–100 risk score is calculated for every physician (NPI) and supplier.")
rs.italic = True; rs.font.size = Pt(9.5)

# Formula
heading("1. Scoring formula")
f = doc.add_paragraph()
fr = f.add_run("risk_score = min( Σ points for each fraud rule that fired  +  min(physician_flags × 5, 20),  100 )")
fr.font.name = "Consolas"; fr.font.size = Pt(10); fr.bold = True
b = doc.add_paragraph()
b.add_run("Each rule contributes its fixed points once if it fires (regardless of how many claims trigger it). "
          "Risk bands: ").font.size = Pt(9.5)
rb = b.add_run("Critical > 80   High > 60   Medium > 30   Low ≤ 30"); rb.bold = True; rb.font.size = Pt(9.5)


def make_table(headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.width = Inches(widths[i])
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9)
            if i == 1:
                run.bold = True
    return tbl


heading("2. Points per fraud pattern (rule weights + trigger thresholds)")
make_table(
    ["Fraud pattern", "Points", "Threshold that triggers it"],
    [
        ["OIG LEIE hit", "+35", "Supplier on the OIG exclusion list (any match)"],
        ["Cross-NPI supplier", "+30", "Supplier bills under > 3 distinct physician NPIs"],
        ["Volume spike", "+25", "Last-30-day claim rate > 2.0x the prior 60-day baseline"],
        ["Duplicate billing", "+20", "Same patient + service + date billed by > 1 supplier"],
        ["Patient identity reuse", "+20", "Same patient billed under >= 3 distinct NPIs"],
        ["Upcoding", "+20", "Claim amount > 3.0x category median and > $500"],
        ["Geographic anomaly", "+15", "Patient > 150 miles from the practice"],
        ["Abnormal hospice duration", "+15", "Hospice enrollment span > 180 days"],
        ["Unbundling", "+15", ">= 3 distinct CPT codes - same patient / date / NPI / supplier"],
        ["New high-value supplier", "+10", "New supplier (first seen <= 30 days) with a claim > $500"],
    ],
    [1.9, 0.7, 4.2],
)

heading("3. What a physician action adds")
make_table(
    ["Physician action", "Points", "Notes"],
    [
        ["Flag Supplier", "+5", "Counts toward the score"],
        ["Unknown Patient", "+5", "Counts toward the score"],
        ["Did Not Order", "0", "Fires a live alert + counts as a denial (not yet scored)"],
        ["Dispute", "0", "Recorded for review; does not change the score"],
        ["Confirm", "0", "Recorded for review; does not change the score"],
    ],
    [1.9, 0.7, 4.2],
)
note = doc.add_paragraph()
nr = note.add_run("Physician contribution is capped at +20 total (the 5th+ counted flag adds nothing). "
                  "Supplier scores exclude Volume spike and Geographic anomaly (physician-level signals). "
                  "All weights and thresholds are configurable in config.py.")
nr.italic = True; nr.font.size = Pt(8.5)

out = r"D:\Mediclaim\docs\ClaimLens_Scoring.docx"
doc.save(out)
print("saved", out)
